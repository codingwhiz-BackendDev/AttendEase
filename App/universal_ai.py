"""
Universal AI Teaching System - The Most Comprehensive Educational AI
Handles ALL subjects and learning types:
- STEM (Science, Technology, Engineering, Math)
- Humanities (History, Literature, Philosophy, Arts)
- Social Sciences (Economics, Psychology, Sociology)
- Business (Management, Finance, Marketing)
- Languages (Multiple language support)
- Creative Writing
- Research Skills
- Critical Thinking
- And much more...
"""

import os
import json
import re
import logging
from datetime import datetime
from typing import List, Dict, Tuple, Optional
import numpy as np
from django.conf import settings
import google.generativeai as genai
from sentence_transformers import SentenceTransformer
import faiss
import sympy as sp
import pypdf
import docx
from io import BytesIO

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Deep document understanding and extraction"""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF with structure preservation"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                text = ""
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text += f"\n\n=== Page {page_num + 1} ===\n{page_text}"
                
                return text
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from DOCX with structure preservation"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n\n"
            
            # Extract tables
            for table in doc.tables:
                text += "\n=== Table ===\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text += row_text + "\n"
                text += "\n"
            
            return text
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            return ""
    
    @staticmethod
    def extract_key_concepts(text: str) -> List[Dict]:
        """Extract key concepts, definitions, and formulas from text"""
        concepts = []
        
        # Extract definitions (common patterns)
        definition_patterns = [
            r'([A-Z][a-zA-Z\s]+)\s+(?:is|refers to|means|can be defined as)\s+([^.]+)\.',
            r'([A-Z][a-zA-Z\s]+):\s+([^.]+)\.',
            r'Definition:\s+([^.]+)\.?\s*([A-Z][a-zA-Z\s]+)'
        ]
        
        for pattern in definition_patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                concepts.append({
                    'type': 'definition',
                    'term': match.group(1).strip(),
                    'definition': match.group(2).strip() if len(match.groups()) > 1 else match.group(1).strip()
                })
        
        # Extract formulas/mathematical expressions
        formula_patterns = [
            r'([A-Z][a-z]+\s*[=≤≥≠<>]+[^.]+\.?)',
            r'\\[([^\\]+)\\]',  # LaTeX formulas
            r'\$([^$]+)\$'  # Inline LaTeX
        ]
        
        for pattern in formula_patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                concepts.append({
                    'type': 'formula',
                    'expression': match.group(1).strip()
                })
        
        # Extract important terms (capitalized words in context)
        term_pattern = r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\b'
        matches = re.findall(term_pattern, text)
        term_counts = {}
        for term in matches:
            if len(term) > 3:  # Filter short words
                term_counts[term] = term_counts.get(term, 0) + 1
        
        # Most frequent terms are likely key concepts
        for term, count in sorted(term_counts.items(), key=lambda x: x[1], reverse=True)[:20]:
            if count >= 2:  # Appears at least twice
                concepts.append({
                    'type': 'key_term',
                    'term': term,
                    'frequency': count
                })
        
        return concepts
    
    @staticmethod
    def extract_sections(text: str) -> List[Dict]:
        """Extract document sections and headings"""
        sections = []
        lines = text.split('\n')
        current_section = None
        
        for line in lines:
            # Detect headings (all caps, or ending with colon, or numbered)
            if line.strip().isupper() and len(line.strip()) < 100:
                if current_section:
                    sections.append(current_section)
                current_section = {
                    'heading': line.strip(),
                    'content': ''
                }
            elif re.match(r'^\d+\.?\s+[A-Z]', line.strip()):
                if current_section:
                    sections.append(current_section)
                current_section = {
                    'heading': line.strip(),
                    'content': ''
                }
            elif line.strip().endswith(':') and len(line.strip()) < 100:
                if current_section:
                    sections.append(current_section)
                current_section = {
                    'heading': line.strip(),
                    'content': ''
                }
            elif current_section:
                current_section['content'] += line + '\n'
        
        if current_section:
            sections.append(current_section)
        
        return sections
    
    @staticmethod
    def generate_document_summary(text: str) -> Dict:
        """Generate a comprehensive summary of the document"""
        concepts = DocumentProcessor.extract_key_concepts(text)
        sections = DocumentProcessor.extract_sections(text)
        
        key_terms = [c.get('term') for c in concepts if c.get('type') == 'key_term'][:10]
        definitions = [c for c in concepts if c.get('type') == 'definition'][:5]
        formulas = [c.get('expression', '') for c in concepts if c.get('type') == 'formula'][:5]
        
        return {
            'total_concepts': len(concepts),
            'key_terms': key_terms,
            'definitions': definitions,
            'formulas': formulas,
            'sections': [s.get('heading', '') for s in sections],
            'sections_count': len(sections)
        }
    
    @staticmethod
    def process_material(material) -> str:
        """Process a course material and extract its content"""
        if material.content and material.content.strip():
            return material.content
        
        if material.file:
            file_path = material.file.path
            if file_path.endswith('.pdf'):
                return DocumentProcessor.extract_text_from_pdf(file_path)
            elif file_path.endswith('.docx'):
                return DocumentProcessor.extract_text_from_docx(file_path)
        
        return ""


class SubjectClassifier:
    """Classify subjects and domains for specialized teaching"""
    
    SUBJECT_DOMAINS = {
        'mathematics': ['math', 'calculus', 'algebra', 'geometry', 'statistics', 'probability', 'linear algebra'],
        'physics': ['physics', 'mechanics', 'thermodynamics', 'electromagnetism', 'quantum', 'optics'],
        'chemistry': ['chemistry', 'organic', 'induction', 'inorganic', 'biochemistry', 'chemical'],
        'biology': ['biology', 'genetics', 'ecology', 'microbiology', 'anatomy', 'physiology'],
        'computer_science': ['programming', 'coding', 'algorithm', 'data structure', 'software', 'ai', 'machine learning'],
        'engineering': ['engineering', 'civil', 'mechanical', 'electrical', 'chemical', 'software'],
        'economics': ['economics', 'microeconomics', 'macroeconomics', 'finance', 'market'],
        'business': ['business', 'management', 'marketing', 'entrepreneurship', 'strategy'],
        'history': ['history', 'historical', 'ancient', 'modern', 'war', 'civilization'],
        'literature': ['literature', 'novel', 'poetry', 'writing', 'author', 'book'],
        'philosophy': ['philosophy', 'ethics', 'logic', 'metaphysics', 'epistemology'],
        'psychology': ['psychology', 'cognitive', 'behavioral', 'mental', 'therapy'],
        'sociology': ['sociology', 'social', 'society', 'culture', 'community'],
        'political_science': ['politics', 'government', 'democracy', 'policy', 'election'],
        'law': ['law', 'legal', 'constitution', 'rights', 'justice'],
        'medicine': ['medicine', 'medical', 'health', 'disease', 'treatment'],
        'environmental': ['environment', 'climate', 'sustainability', 'ecology', 'green'],
        'arts': ['art', 'design', 'music', 'painting', 'sculpture', 'architecture'],
        'languages': ['language', 'grammar', 'vocabulary', 'translation', 'linguistics']
    }
    
    @classmethod
    def classify_subject(cls, text: str) -> Tuple[str, str]:
        """Classify the subject and domain from text"""
        text_lower = text.lower()
        
        for domain, keywords in cls.SUBJECT_DOMAINS.items():
            for keyword in keywords:
                if keyword in text_lower:
                    return domain, keyword
        
        return 'general', 'general'


class LearningStyleDetector:
    """Detect and adapt to different learning styles"""
    
    LEARNING_STYLES = {
        'visual': ['see', 'look', 'watch', 'diagram', 'chart', 'graph', 'picture', 'visual'],
        'auditory': ['hear', 'listen', 'sound', 'audio', 'speak', 'discuss', 'explain'],
        'kinesthetic': ['do', 'practice', 'hands-on', 'try', 'experiment', 'build', 'make'],
        'reading_writing': ['read', 'write', 'text', 'note', 'book', 'article', 'document']
    }
    
    @classmethod
    def detect_style(cls, query: str, history: List[Dict] = None) -> Dict:
        """Detect student's preferred learning style"""
        query_lower = query.lower()
        scores = {style: 0 for style in cls.LEARNING_STYLES}
        
        # Score based on query
        for style, keywords in cls.LEARNING_STYLES.items():
            for keyword in keywords:
                if keyword in query_lower:
                    scores[style] += 1
        
        # Consider history if available
        if history:
            for interaction in history[-10:]:
                interaction_text = interaction.get('query', '').lower()
                for style, keywords in cls.LEARNING_STYLES.items():
                    for keyword in keywords:
                        if keyword in interaction_text:
                            scores[style] += 0.5
        
        # Determine dominant style
        max_score = max(scores.values())
        if max_score == 0:
            return {'style': 'balanced', 'scores': scores}
        
        dominant = [s for s, score in scores.items() if score == max_score][0]
        return {'style': dominant, 'scores': scores}


class BloomTaxonomyAdapter:
    """Adapt teaching to Bloom's Taxonomy levels"""
    
    LEVELS = {
        'remember': ['define', 'list', 'identify', 'name', 'recall', 'what is'],
        'understand': ['explain', 'describe', 'discuss', 'summarize', 'interpret'],
        'apply': ['apply', 'use', 'implement', 'solve', 'calculate'],
        'analyze': ['analyze', 'compare', 'contrast', 'differentiate', 'examine'],
        'evaluate': ['evaluate', 'assess', 'critique', 'judge', 'justify'],
        'create': ['create', 'design', 'construct', 'develop', 'formulate']
    }
    
    @classmethod
    def detect_level(cls, query: str) -> str:
        """Detect Bloom's taxonomy level from query"""
        query_lower = query.lower()
        
        for level, verbs in cls.LEVELS.items():
            for verb in verbs:
                if verb in query_lower:
                    return level
        
        return 'understand'  # Default level
    
    @classmethod
    def get_teaching_approach(cls, level: str) -> str:
        """Get appropriate teaching approach for Bloom's level"""
        approaches = {
            'remember': 'Use repetition, mnemonics, and recall exercises',
            'understand': 'Use explanations, analogies, and examples',
            'apply': 'Use practice problems, case studies, and applications',
            'analyze': 'Use comparisons, breakdowns, and critical thinking',
            'evaluate': 'Use critiques, assessments, and justification exercises',
            'create': 'Use projects, design challenges, and creative tasks'
        }
        return approaches.get(level, approaches['understand'])


class UniversalTeachingStrategies:
    """Subject-specific and domain-specific teaching strategies"""
    
    @staticmethod
    def stem_teaching(concept: str, domain: str) -> str:
        """Teaching strategy for STEM subjects"""
        strategies = {
            'mathematics': """
                For mathematical concepts:
                1. Start with the problem this math solves
                2. Show the formula with LaTeX notation
                3. Derive it step-by-step
                4. Provide numerical examples
                5. Show real-world applications
                6. Give practice problems
            """,
            'physics': """
                For physics concepts:
                1. Describe the physical phenomenon
                2. Show the governing equations
                3. Explain the physical intuition
                4. Provide real-world examples
                5. Include thought experiments
                6. Suggest laboratory verification
            """,
            'computer_science': """
                For computer science concepts:
                1. Explain the problem being solved
                2. Show algorithm or code
                3. Walk through the logic step-by-step
                4. Discuss time/space complexity
                5. Provide optimization alternatives
                6. Give coding exercises
            """,
            'chemistry': """
                For chemistry concepts:
                1. Describe the chemical process
                2. Show molecular structures
                3. Explain the reaction mechanism
                4. Provide real-world applications
                5. Discuss safety considerations
                6. Suggest laboratory experiments
            """
        }
        return strategies.get(domain, strategies['mathematics'])
    
    @staticmethod
    def humanities_teaching(concept: str, domain: str) -> str:
        """Teaching strategy for humanities"""
        strategies = {
            'history': """
                For historical concepts:
                1. Set the historical context
                2. Describe key events and figures
                3. Explain causes and consequences
                4. Connect to present-day relevance
                5. Discuss different perspectives
                6. Suggest primary source analysis
            """,
            'literature': """
                For literary analysis:
                1. Summarize the work's context
                2. Analyze themes and motifs
                3. Discuss character development
                4. Examine literary devices
                5. Connect to broader themes
                6. Encourage critical interpretation
            """,
            'philosophy': """
                For philosophical concepts:
                1. State the philosophical question
                2. Present key arguments
                3. Discuss counterarguments
                4. Examine implications
                5. Connect to other philosophers
                6. Encourage critical reflection
            """
        }
        return strategies.get(domain, strategies['history'])
    
    @staticmethod
    def social_sciences_teaching(concept: str, domain: str) -> str:
        """Teaching strategy for social sciences"""
        strategies = {
            'economics': """
                For economic concepts:
                1. Define the economic principle
                2. Show the economic model
                3. Provide real-world examples
                4. Discuss policy implications
                5. Examine different schools of thought
                6. Suggest data analysis exercises
            """,
            'psychology': """
                For psychological concepts:
                1. Define the psychological phenomenon
                2. Explain underlying mechanisms
                3. Provide experimental evidence
                4. Discuss real-world applications
                5. Examine individual differences
                6. Suggest observational exercises
            """,
            'sociology': """
                For sociological concepts:
                1. Define the social phenomenon
                2. Explain sociological theories
                3. Provide empirical evidence
                4. Discuss social implications
                5. Examine cross-cultural variations
                6. Suggest research methods
            """
        }
        return strategies.get(domain, strategies['economics'])
    
    @staticmethod
    def creative_teaching(concept: str, domain: str) -> str:
        """Teaching strategy for creative subjects"""
        strategies = {
            'arts': """
                For artistic concepts:
                1. Describe the artistic principle
                2. Show visual examples
                3. Explain techniques
                4. Discuss historical context
                5. Provide practice exercises
                6. Encourage creative exploration
            """,
            'literature': """
                For creative writing:
                1. Analyze the writing technique
                2. Show examples from literature
                3. Explain the craft elements
                4. Provide writing exercises
                5. Discuss style and voice
                6. Encourage original work
            """
        }
        return strategies.get(domain, strategies['arts'])


class CrossDomainConnector:
    """Connect concepts across different domains"""
    
    CONNECTIONS = {
        'mathematics': {
            'physics': 'Calculus describes motion and change in physical systems',
            'computer_science': 'Linear algebra underpins machine learning algorithms',
            'economics': 'Statistics models market behavior and probability',
            'engineering': 'Differential equations model engineering systems'
        },
        'history': {
            'literature': 'Historical context shapes literary themes and settings',
            'philosophy': 'Philosophical movements drive historical change',
            'politics': 'Historical events form political institutions',
            'economics': 'Economic forces shape historical developments'
        },
        'science': {
            'philosophy': 'Scientific method has philosophical foundations',
            'mathematics': 'Mathematics provides the language of science',
            'technology': 'Scientific discoveries drive technological innovation',
            'ethics': 'Scientific advances raise ethical questions'
        }
    }
    
    @classmethod
    def find_connections(cls, domain: str, concept: str, mastered_domains: set) -> List[str]:
        """Find connections to other domains the student has mastered"""
        connections = []
        
        if domain in cls.CONNECTIONS:
            for other_domain, description in cls.CONNECTIONS[domain].items():
                if other_domain in mastered_domains:
                    connections.append(f"Connection to {other_domain}: {description}")
        
        return connections


class CriticalThinkingPromoter:
    """Promote critical thinking skills"""
    
    @staticmethod
    def generate_critical_questions(concept: str, domain: str) -> List[str]:
        """Generate critical thinking questions"""
        base_questions = [
            f"What assumptions underlie the concept of {concept}?",
            f"How would {concept} change if we modified a key variable?",
            f"What are the limitations of {concept}?",
            f"How does {concept} compare to alternative approaches?",
            f"What evidence supports {concept}?",
            f"What are the counterarguments to {concept}?",
            f"How would you apply {concept} in a novel situation?",
            f"What are the ethical implications of {concept}?"
        ]
        
        domain_specific = {
            'mathematics': [
                f"Can {concept} be generalized to other contexts?",
                f"What happens at the boundaries or edge cases of {concept}?"
            ],
            'history': [
                f"How might {concept} be interpreted differently?",
                f"What primary sources support our understanding of {concept}?"
            ],
            'science': [
                f"What experiments could test {concept}?",
                f"Is {concept} falsifiable? How would we know if it's wrong?"
            ],
            'philosophy': [
                f"What are the logical consequences of {concept}?",
                f"Does {concept} lead to any paradoxes?"
            ]
        }
        
        questions = base_questions
        if domain in domain_specific:
            questions.extend(domain_specific[domain])
        
        return questions[:6]  # Return top 6


class ResearchAssistant:
    """Help with research skills and academic writing"""
    
    @staticmethod
    def suggest_research_approach(topic: str, domain: str) -> Dict:
        """Suggest research methodology for a topic"""
        methodologies = {
            'mathematics': ['Proof', 'Numerical simulation', 'Theoretical derivation'],
            'physics': ['Experiment', 'Simulation', 'Theoretical analysis'],
            'computer_science': ['Algorithm design', 'Implementation', 'Performance analysis'],
            'history': ['Primary source analysis', 'Comparative analysis', 'Archival research'],
            'social_sciences': ['Survey', 'Case study', 'Statistical analysis', 'Qualitative interviews'],
            'literature': ['Textual analysis', 'Comparative study', 'Historical context'],
            'philosophy': ['Logical analysis', 'Comparative philosophy', 'Conceptual analysis']
        }
        
        methods = methodologies.get(domain, ['Literature review', 'Analysis', 'Synthesis'])
        
        return {
            'topic': topic,
            'domain': domain,
            'suggested_methods': methods,
            'research_questions': [
                f"What is the current understanding of {topic}?",
                f"What are the gaps in knowledge about {topic}?",
                f"How can {topic} be advanced?"
            ],
            'sources': [
                "Academic journals",
                "Books and monographs",
                "Conference proceedings",
                "Theses and dissertations"
            ]
        }


class DebateFacilitator:
    """Facilitate discussions and debates"""
    
    @staticmethod
    def generate_debate_positions(topic: str, domain: str) -> Dict:
        """Generate opposing positions for debate"""
        return {
            'topic': topic,
            'position_a': f"Arguments in favor of {topic}",
            'position_b': f"Arguments against {topic}",
            'key_questions': [
                f"What are the strongest arguments for {topic}?",
                f"What are the strongest arguments against {topic}?",
                f"What evidence supports each position?",
                f"What are the implications of each position?"
            ],
            'evaluation_criteria': [
                "Logical consistency",
                "Evidence quality",
                "Counterargument handling",
                "Practical implications"
            ]
        }


class ProjectBasedLearningGuide:
    """Guide project-based learning"""
    
    @staticmethod
    def suggest_project(concept: str, domain: str, difficulty: str = 'medium') -> Dict:
        """Suggest a project to learn the concept"""
        project_templates = {
            'mathematics': {
                'easy': 'Create a calculator or visualization tool',
                'medium': 'Build a mathematical model of a real-world system',
                'hard': 'Develop a new proof or mathematical framework'
            },
            'computer_science': {
                'easy': 'Build a simple application',
                'medium': 'Create a full-stack project',
                'hard': 'Contribute to open-source or research project'
            },
            'science': {
                'easy': 'Design a simple experiment',
                'medium': 'Conduct a full research investigation',
                'hard': 'Publish original research'
            },
            'humanities': {
                'easy': 'Write an analytical essay',
                'medium': 'Create a multimedia presentation',
                'hard': 'Produce original research or creative work'
            }
        }
        
        templates = project_templates.get(domain, project_templates['mathematics'])
        
        return {
            'concept': concept,
            'domain': domain,
            'difficulty': difficulty,
            'project_idea': templates.get(difficulty, templates['medium']),
            'deliverables': [
                "Project proposal",
                "Progress reports",
                "Final presentation",
                "Documentation"
            ],
            'timeline': f"4-8 weeks for {difficulty} project",
            'learning_objectives': [
                f"Deep understanding of {concept}",
                "Practical application skills",
                "Problem-solving ability",
                "Communication skills"
            ]
        }


class UniversalSmartTutor:
    """
    The Ultimate Universal AI Tutor - Handles ALL Subjects and Learning Types
    Most comprehensive educational AI system ever built
    """
    
    def __init__(self, course_code: str, student_id: str):
        self.course_code = course_code
        self.student_id = student_id
        
        # Initialize all components
        self.subject_classifier = SubjectClassifier()
        self.learning_style_detector = LearningStyleDetector()
        self.bloom_adapter = BloomTaxonomyAdapter()
        self.teaching_strategies = UniversalTeachingStrategies()
        self.cross_domain_connector = CrossDomainConnector()
        self.critical_thinking = CriticalThinkingPromoter()
        self.research_assistant = ResearchAssistant()
        self.debate_facilitator = DebateFacilitator()
        self.project_guide = ProjectBasedLearningGuide()
        self.document_processor = DocumentProcessor()
        
        # Vector store and chunker
        from App.smart_ai import DocumentChunker, VectorStore
        self.chunker = DocumentChunker()
        
        # Force CPU device to avoid meta tensor issues
        import torch
        torch.device('cpu')
        self.vector_store = VectorStore()
        
        # Also ensure the vector store uses CPU
        import torch
        self.vector_store.device = 'cpu'
        
        # Math solver
        from App.smart_ai import MathSolver
        self.math_solver = MathSolver()
        
        # Learning memory
        self.topics_mastered = set()
        self.domains_mastered = set()
        self.topics_struggling = set()
        self.conversation_history = []
        self.detected_learning_style = 'balanced'
        
        # Document understanding cache
        self.document_summaries = {}
        
        # Quota management
        self.quota_retries = 0
        self.max_quota_retries = 3
        
        # Initialize Gemini
        api_key = settings.GEMINI_API_KEY
        model_name = settings.GEMINI_MODEL
        
        if not api_key or api_key == 'your_gemini_api_key_here':
            raise ValueError("GEMINI_API_KEY not configured")
        
        # Model management for quota handling
        self.current_model = None
        self.available_models = []
        self.safe_models = []  # Models that are not unavailable
        self.model_index = 0
        self.last_quota_error = None
        
        # Try to list available models and use a working one
        try:
            models = genai.list_models()
            model_list = list(models)
            self.available_models = [m.name for m in model_list]
            
            # Try different models in order of preference (only available models)
            # Define unavailable models that should be avoided
            self.unavailable_models = ['models/gemini-2.5-pro', 'models/gemini-2.5-flash']
            
            # Filter available models to only include safe ones
            self.safe_models = [m for m in self.available_models if m not in self.unavailable_models]
            
            if not self.safe_models:
                raise ValueError("No available models (all models are unavailable)")
            
            # Define preferred models in order
            preferred_models = [
                'models/gemini-3.7-flash',
                'models/gemini-3.6-flash', 
                'models/gemini-3.5-flash',
                'models/gemini-flash-latest'
            ]
            
            # Find the first preferred model that's available
            for model in preferred_models:
                if model in self.safe_models:
                    self.current_model = model
                    self.model = genai.GenerativeModel(model)
                    logger.info(f"Using model: {model}")
                    break
            else:
                # Fallback to first safe model
                self.current_model = self.safe_models[0]
                self.model = genai.GenerativeModel(self.current_model)
                logger.info(f"Using fallback model: {self.current_model}")
                
        except Exception as e:
            # Fallback to the configured model
            genai.configure(api_key=api_key)
            self.current_model = model_name
            self.model = genai.GenerativeModel(model_name)
        
        # Load course materials
        self._load_course_materials()
    
    def _load_course_materials(self):
        """Load and deeply understand course materials including student uploads"""
        try:
            from App.models import Course, CourseMaterial, StudentDocument
            course = Course.objects.get(course_code=self.course_code)
            materials = CourseMaterial.objects.filter(course=course)
            
            # Also load student documents for this course
            try:
                student_docs = StudentDocument.objects.filter(course_code=self.course_code, student__id=int(self.student_id))
            except:
                student_docs = StudentDocument.objects.none()
            
            all_chunks = []
            
            # Process lecturer materials
            for material in materials:
                # Process material to extract content
                content = self.document_processor.process_material(material)
                
                if content:
                    # Generate document summary
                    summary = self.document_processor.generate_document_summary(content)
                    self.document_summaries[material.id] = {
                        'title': material.title,
                        'type': material.material_type,
                        'summary': summary,
                        'source': 'lecturer'
                    }
                    
                    # Chunk the content with enhanced metadata
                    chunks = self.chunker.chunk_text(
                        content,
                        metadata={
                            'material_id': material.id,
                            'title': material.title,
                            'type': material.material_type,
                            'key_concepts': summary['key_terms'][:5],
                            'source': 'lecturer'
                        }
                    )
                    all_chunks.extend(chunks)
            
            # Process student documents
            for doc in student_docs:
                try:
                    content = ""
                    if doc.file:
                        if doc.file_type == 'pdf':
                            import pypdf
                            pdf_reader = pypdf.PdfReader(doc.file.path)
                            for page in pdf_reader.pages:
                                content += page.extract_text()
                        elif doc.file_type == 'docx':
                            from docx import Document
                            doc_obj = Document(doc.file.path)
                            content = "\n".join([para.text for para in doc_obj.paragraphs])
                        elif doc.file_type == 'txt':
                            content = doc.file.read().decode('utf-8')
                    
                    if content.strip():
                        # Generate simple summary
                        summary = content[:500]
                        self.document_summaries[f"student_{doc.id}"] = {
                            'title': doc.title,
                            'type': doc.file_type,
                            'summary': summary,
                            'source': 'student'
                        }
                        
                        # Chunk the content
                        chunks = self.chunker.chunk_text(
                            content,
                            metadata={
                                'material_id': f"student_{doc.id}",
                                'title': doc.title,
                                'type': doc.file_type,
                                'source': 'student'
                            }
                        )
                        all_chunks.extend(chunks)
                except Exception as e:
                    logger.error(f"Error processing student document {doc.id}: {str(e)}")
            
            if all_chunks:
                self.vector_store.add_chunks(all_chunks)
                logger.info(f"Loaded {len(all_chunks)} chunks from {len(materials)} lecturer materials and {len(student_docs)} student documents for {self.course_code}")
                logger.info(f"Document summaries: {len(self.document_summaries)} materials processed")
            
        except Exception as e:
            logger.error(f"Error loading materials: {str(e)}")
    
    def _retrieve_relevant_context(self, query: str, top_k: int = 5) -> List[Dict]:
        """Retrieve most relevant document chunks"""
        from App.smart_ai import VectorStore
        results = self.vector_store.search(query, top_k)
        
        context = []
        for text, metadata, score in results:
            context.append({
                'text': text,
                'metadata': metadata,
                'relevance_score': score
            })
        
        return context
    
    def _generate_follow_up(self, query: str, answer: str, domain: str) -> str:
        """Generate a relevant follow-up question to test understanding"""
        try:
            # Simple follow-up based on the topic
            follow_up_prompt = f"""Based on this Q&A:
Student asked: {query}
You answered: {answer[:200]}...

Generate ONE follow-up question to test if the student understood. Keep it simple and relevant."""
            
            response = self.model.generate_content(follow_up_prompt)
            return response.text.strip()
        except:
            return "Does this explanation make sense? Would you like me to try a different approach?"
    
    def _extract_citations_from_context(self, context: List[Dict]) -> List[str]:
        """Extract source citations from the context"""
        citations = []
        for ctx in context[:3]:  # Top 3 most relevant
            metadata = ctx.get('metadata', {})
            title = metadata.get('title', 'Unknown')
            source = metadata.get('source', 'lecturer')
            citations.append(f"{source}: {title}")
        return citations
    
    def _detect_intent(self, query: str) -> str:
        """Detect the student's intent to route to appropriate functionality"""
        query_lower = query.lower()
        
        # Quiz generation intent
        quiz_keywords = ['quiz', 'test', 'exam', 'mcq', 'multiple choice', 'question me', 'give me questions']
        if any(keyword in query_lower for keyword in quiz_keywords):
            return 'quiz'
        
        # Flashcard intent
        flashcard_keywords = ['flashcard', 'flash card', 'flip card', 'term', 'definition']
        if any(keyword in query_lower for keyword in flashcard_keywords):
            return 'flashcards'
        
        # Summary intent
        summary_keywords = ['summarize', 'summary', 'summarize this', 'key points', 'main points', 'overview']
        if any(keyword in query_lower for keyword in summary_keywords):
            return 'summarize'
        
        # Simple explanation intent
        simple_keywords = ['explain simply', 'explain in simple terms', 'dumb it down', 'like i\'m 5', 'easy explanation']
        if any(keyword in query_lower for keyword in simple_keywords):
            return 'explain_simply'
        
        # Deep dive intent
        deep_keywords = ['deep dive', 'detailed', 'in depth', 'thorough', 'comprehensive', 'advanced']
        if any(keyword in query_lower for keyword in deep_keywords):
            return 'deep_dive'
        
        # Exam prep intent
        exam_keywords = ['exam prep', 'exam preparation', 'study for exam', 'past questions', 'revision']
        if any(keyword in query_lower for keyword in exam_keywords):
            return 'exam_prep'
        
        # Practice intent
        practice_keywords = ['practice', 'give me a question', 'ask me', 'test my knowledge']
        if any(keyword in query_lower for keyword in practice_keywords):
            return 'practice'
        
        # Default to general chat
        return 'chat'
    
    def _get_learning_mode_instruction(self, intent: str) -> str:
        """Get specific instructions based on detected intent"""
        instructions = {
            'quiz': "Generate a quiz question with multiple choice options from the course materials.",
            'flashcards': "Create a flashcard with a question/term on the front and answer/definition on the back.",
            'summarize': "Provide a concise summary of the key points from the relevant materials.",
            'explain_simply': "Explain this concept in the simplest possible terms, as if teaching someone who has never seen it before.",
            'deep_dive': "Provide a comprehensive, detailed explanation covering all aspects of this topic.",
            'exam_prep': "Focus on exam-relevant information, common question patterns, and what students should know for tests.",
            'practice': "Ask a practice question to test understanding, then be ready to evaluate their answer.",
            'chat': "Answer the student's question helpfully and naturally, like a patient tutor."
        }
        return instructions.get(intent, instructions['chat'])
    
    def _build_universal_teaching_prompt(self, query: str, context: List[Dict], domain: str, 
                                           learning_style: str, bloom_level: str) -> str:
        """Build the most comprehensive teaching prompt for ANY subject"""
        
        # Detect intent and get appropriate instructions
        intent = self._detect_intent(query)
        mode_instruction = self._get_learning_mode_instruction(intent)
        
        # Simple, natural teaching approach - MilliaAi the friendly tutor
        prompt = f"""You are MilliaAi, a patient, friendly AI tutor who helps students understand any subject. You teach like the smartest, most patient tutor in the world - you can help even the most confused student understand complex topics.

Your teaching style:
- Be conversational and friendly, like a helpful study buddy
- Explain things simply, step by step
- Use real examples and analogies that make sense
- If something is confusing, break it down into smaller pieces
- Be encouraging and positive
- Ask if they understand before moving on
- If they're stuck, try a different explanation approach

Current situation:
- Subject: {domain}
- Course: {self.course_code}
- Learning mode: {intent}
- Specific instruction: {mode_instruction}

Student's question: "{query}"

Available course materials (if any):
"""
        
        # Add relevant context from materials with source information
        if context:
            for i, ctx in enumerate(context[:3], 1):
                source = ctx['metadata'].get('source', 'lecturer')
                title = ctx['metadata'].get('title', 'Unknown')
                prompt += f"\nMaterial {i} ({source} - {title}): {ctx['text'][:500]}...\n"
        else:
            prompt += "\nNo specific course materials available for this question.\n"
        
        prompt += f"""

Answer the student's question helpfully and naturally following the learning mode instruction. Don't mention teaching strategies or learning styles - just explain the topic clearly and patiently, like a good tutor would.
"""
        
        return prompt
    
    def teach(self, query: str) -> Dict:
        """
        Universal teaching method - works for ANY subject
        """
        # Handle simple greetings with friendly responses
        greetings = ['hi', 'hello', 'hey', 'good morning', 'good afternoon', 'good evening', 'thanks', 'thank you']
        if query.lower().strip() in greetings or query.lower().strip().startswith(tuple(greetings)):
            return {
                'answer': f"Hello! 👋 I'm your AI study tutor for {self.course_code}. How can I help you today? Feel free to ask questions about your course materials, request explanations, or generate practice problems!",
                'domain': 'general',
                'learning_style': 'balanced',
                'bloom_level': 'remember',
                'confidence': 1.0,
                'citations': [],
                'math_solution': None,
                'visual_aids': [],
                'practice_questions': [],
                'source_materials': []
            }
        
        # Classify subject and domain
        domain, keyword = self.subject_classifier.classify_subject(query)
        
        # Detect learning style
        style_info = self.learning_style_detector.detect_style(query, self.conversation_history)
        learning_style = style_info['style']
        self.detected_learning_style = learning_style
        
        # Detect Bloom's level
        bloom_level = self.bloom_adapter.detect_level(query)
        
        # Handle math specially
        if domain == 'mathematics' and any(c in query for c in ['=', '+', '-', '*', '/', '^']):
            math_result = self.math_solver.solve_equation(query)
            if math_result['success']:
                return self._format_math_response(math_result, query)
        
        # Retrieve relevant context
        context = self._retrieve_relevant_context(query)
        
        # Build universal teaching prompt
        prompt = self._build_universal_teaching_prompt(query, context, domain, learning_style, bloom_level)
        
        try:
            # Generate response with quota handling
            response = self.model.generate_content(prompt)
            answer = response.text
            
            # Reset quota retries on success
            self.quota_retries = 0
            
            # Post-process
            answer = self._enhance_response(answer, domain)
            
            # Extract citations
            citations = self._extract_citations(answer)
            
            # Calculate confidence
            confidence = self._calculate_confidence(context)
            
            # Update learning memory
            self.conversation_history.append({
                'timestamp': datetime.now().isoformat(),
                'query': query,
                'response': answer,
                'domain': domain,
                'learning_style': learning_style,
                'bloom_level': bloom_level,
                'intent': intent
            })
            
            # Generate follow-up question for engagement
            follow_up = self._generate_follow_up(query, answer, domain)
            
            # Extract citations from context
            citations = self._extract_citations_from_context(context) if context else []
            
            # Update domain mastery
            if confidence > 0.8:
                self.domains_mastered.add(domain)
            
            return {
                'answer': answer,
                'citations': citations,
                'context_used': len(context),
                'confidence': confidence,
                'domain': domain,
                'keyword': keyword,
                'learning_style': learning_style,
                'bloom_level': bloom_level,
                'intent': intent,
                'follow_up_question': follow_up,
                'learning_state': {
                    'mastered_topics': list(self.topics_mastered),
                    'mastered_domains': list(self.domains_mastered),
                    'struggling_topics': list(self.topics_struggling)
                }
            }
            
        except Exception as e:
            error_msg = str(e)
            
            # Handle quota errors specifically
            if '429' in error_msg or 'quota' in error_msg.lower():
                logger.warning(f"Quota error: {error_msg}")
                
                # Try to switch to a different model
                if self.quota_retries < self.max_quota_retries:
                    self.quota_retries += 1
                    logger.info(f"Quota exceeded, trying alternative model (attempt {self.quota_retries})")
                    
                    # Try next available model
                    if self.model_index < len(self.safe_models) - 1:
                        self.model_index += 1
                        self.current_model = self.safe_models[self.model_index]
                        self.model = genai.GenerativeModel(self.current_model)
                        logger.info(f"Switched to model: {self.current_model}")
                        
                        # Retry the request
                        return self.teach(query)
                    else:
                        self.model_index = 0  # Reset for next time
                        return {
                            'error': 'quota_exceeded',
                            'answer': "⚠️ AI quota exceeded. Please try again in a few minutes or upgrade your API plan for higher limits.",
                            'domain': domain,
                            'learning_style': learning_style,
                            'bloom_level': bloom_level,
                            'confidence': 0.0,
                            'citations': [],
                            'context_used': 0
                        }
                else:
                    return {
                        'error': 'quota_exceeded',
                        'answer': "⚠️ AI quota exceeded. Please try again in a few minutes or upgrade your API plan for higher limits.",
                        'domain': domain,
                        'learning_style': learning_style,
                        'bloom_level': bloom_level,
                        'confidence': 0.0,
                        'citations': [],
                        'context_used': 0
                    }
            
            # Handle other errors
            logger.error(f"Universal teaching error: {error_msg}")
            return {
                'error': error_msg,
                'answer': f"⚠️ Error: {error_msg}",
                'domain': domain,
                'learning_style': learning_style,
                'bloom_level': bloom_level,
                'confidence': 0.0,
                'citations': [],
                'context_used': 0
            }
    
    def _format_math_response(self, math_result: Dict, original_query: str) -> Dict:
        """Format mathematical solution response"""
        answer = f"""## Mathematical Solution

**Equation:** {original_query}

### Step-by-Step Solution:
"""
        for step in math_result['steps']:
            answer += f"{step}\n\n"
        
        answer += f"""
### LaTeX Representation:
**Equation:** ${math_result['latex_equation']}$
**Solution:** ${math_result['latex_solution']}$

### Final Answer:
{math_result['solution']}

### Practice:
Try solving a similar problem to reinforce your understanding."""
        
        return {
            'answer': answer,
            'citations': [],
            'context_used': 0,
            'confidence': 0.95,
            'domain': 'mathematics',
            'keyword': 'math',
            'learning_style': 'analytical',
            'bloom_level': 'apply',
            'critical_questions': self.critical_thinking.generate_critical_questions(original_query, 'mathematics'),
            'learning_state': {
                'mastered_topics': list(self.topics_mastered),
                'mastered_domains': list(self.domains_mastered),
                'struggling_topics': list(self.topics_struggling)
            }
        }
    
    def _enhance_response(self, answer: str, domain: str) -> str:
        """Enhance response based on domain"""
        # Ensure LaTeX is properly formatted
        answer = re.sub(r'\$(.*?)\$', r'$\1$', answer)
        
        # Add domain-specific visual cues
        if domain == 'mathematics':
            answer = answer.replace('**Formula:**', '\n📐 **Formula:**')
        elif domain in ['history', 'literature', 'philosophy']:
            answer = answer.replace('**Context:**', '\n📖 **Context:**')
        elif domain in ['physics', 'chemistry', 'biology']:
            answer = answer.replace('**Experiment:**', '\n🔬 **Experiment:**')
        elif domain in ['economics', 'business']:
            answer = answer.replace('**Application:**', '\n💼 **Application:**')
        
        return answer
    
    def _extract_citations(self, response: str) -> List[str]:
        """Extract material citations from response"""
        citations = re.findall(r'\[From: ([^\]]+)\]', response)
        return citations
    
    def _calculate_confidence(self, context: List[Dict]) -> float:
        """Calculate confidence based on context relevance"""
        if not context:
            return 0.5
        
        avg_score = sum(c['relevance_score'] for c in context) / len(context)
        return min(avg_score, 1.0)
    
    def suggest_research(self, topic: str) -> Dict:
        """Suggest research approach for a topic"""
        domain, _ = self.subject_classifier.classify_subject(topic)
        return self.research_assistant.suggest_research_approach(topic, domain)
    
    def generate_debate(self, topic: str) -> Dict:
        """Generate debate positions"""
        domain, _ = self.subject_classifier.classify_subject(topic)
        return self.debate_facilitator.generate_debate_positions(topic, domain)
    
    def suggest_project(self, concept: str, difficulty: str = 'medium') -> Dict:
        """Suggest project-based learning"""
        domain, _ = self.subject_classifier.classify_subject(concept)
        return self.project_guide.suggest_project(concept, domain, difficulty)
    
    def generate_deep_quiz(self, num_questions: int, difficulty: str, topics: str = "") -> Dict:
        """Generate quiz based on DEEP understanding of course materials"""
        # Get document summaries to understand content structure
        material_summaries = list(self.document_summaries.values())
        
        # Build context about available materials
        materials_context = ""
        for summary in material_summaries:
            materials_context += f"\nMaterial: {summary['title']}\n"
            materials_context += f"Type: {summary['type']}\n"
            materials_context += f"Key Concepts: {', '.join(summary['summary']['key_terms'])}\n"
            materials_context += f"Definitions: {len(summary['summary']['definitions'])}\n"
            materials_context += f"Formulas: {len(summary['summary']['formulas'])}\n"
            materials_context += f"Sections: {', '.join(summary['summary']['sections'][:5])}\n"
        
        prompt = f"""Generate {num_questions} multiple-choice questions based on the ACTUAL course materials provided below.

AVAILABLE MATERIALS:
{materials_context if materials_context else "No materials available yet"}

TOPICS TO COVER: {topics if topics else "All key concepts from materials"}
DIFFICULTY: {difficulty}

CRITICAL REQUIREMENTS:
1. Questions MUST be based on the actual content in the materials
2. Use the specific definitions, formulas, and concepts found in the materials
3. Reference the specific material source in explanations
4. Test understanding, not just memorization
5. Include questions that require application of concepts
6. Make options plausible but clearly wrong
7. Provide detailed explanations citing the material

Format each question as JSON:
{{
    "question": "Question text from materials",
    "options": ["A", "B", "C", "D"],
    "correct_answer": 0,
    "explanation": "Detailed explanation with material reference [From: Material Name]",
    "source_material": "Name of material this question comes from",
    "key_concept": "Which concept this tests",
    "difficulty": "easy/medium/hard"
}}

Return the entire quiz as a JSON array of questions.
Questions should be challenging and test deep understanding of the material."""
        
        try:
            response = self.model.generate_content(prompt)
            quiz_json = self._extract_json(response.text)
            return quiz_json if quiz_json else {"error": "Could not parse quiz JSON"}
        except Exception as e:
            logger.error(f"Deep quiz generation error: {str(e)}")
            return {"error": str(e)}
    
    def generate_deep_flashcards(self, num_cards: int, topics: str = "") -> Dict:
        """Generate flashcards based on DEEP understanding of course materials"""
        # Get document summaries to understand content structure
        material_summaries = list(self.document_summaries.values())
        
        # Build context about available materials
        materials_context = ""
        for summary in material_summaries:
            materials_context += f"\nMaterial: {summary['title']}\n"
            materials_context += f"Key Terms: {', '.join(summary['summary']['key_terms'])}\n"
            if summary['summary']['definitions']:
                materials_context += f"Available Definitions: {', '.join([d['term'] for d in summary['summary']['definitions']])}\n"
        
        prompt = f"""Generate {num_cards} flashcards based on the ACTUAL course materials provided below.

AVAILABLE MATERIALS:
{materials_context if materials_context else "No materials available yet"}

TOPICS TO COVER: {topics if topics else "All key concepts from materials"}

CRITICAL REQUIREMENTS:
1. Flashcards MUST be based on actual content in the materials
2. Use the specific definitions and concepts found in the materials
3. Front should be a key term, concept, or question from materials
4. Back should be the exact definition or explanation from materials
5. Include the material source
6. Make terms specific to the course content
7. Include both definitions and concept explanations

Format each flashcard as JSON:
{{
    "front": "Term or question from materials",
    "back": "Definition or answer from materials",
    "source_material": "Name of material",
    "type": "definition/concept/formula/example",
    "key_concept": "Which concept this reinforces"
}}

Return the entire set as a JSON array of flashcards.
Flashcards should test the most important concepts from the materials."""
        
        try:
            response = self.model.generate_content(prompt)
            flashcards_json = self._extract_json(response.text)
            return flashcards_json if flashcards_json else {"error": "Could not parse flashcards JSON"}
        except Exception as e:
            logger.error(f"Deep flashcard generation error: {str(e)}")
            return {"error": str(e)}
    
    def get_material_coverage(self) -> Dict:
        """Get information about what materials are available"""
        return {
            'total_materials': len(self.document_summaries),
            'materials': [
                {
                    'title': s['title'],
                    'type': s['type'],
                    'key_concepts': s['summary']['key_terms'],
                    'has_definitions': len(s['summary']['definitions']) > 0,
                    'has_formulas': len(s['summary']['formulas']) > 0,
                    'sections': s['summary']['sections_count']
                }
                for s in self.document_summaries.values()
            ]
        }
    
    def _extract_json(self, text: str) -> List:
        """Extract JSON array from response"""
        json_match = re.search(r'\[.*\]', text, re.DOTALL)
        if json_match:
            try:
                return json.loads(json_match.group())
            except:
                pass
        return []
    
    def _extract_json_object(self, text: str) -> Dict:
        """Extract JSON object from response"""
        json_match = re.search(r'\{.*\}', text, re.DOTALL)
        if json_match:
            try:
                return json.loads(json_match.group())
            except:
                pass
        return {}


# Singleton tutor instances
_tutor_cache = {}
_current_model = None

def get_universal_tutor(course_code: str, student_id: str) -> UniversalSmartTutor:
    """Get or create a universal smart tutor instance"""
    # TEMPORARILY DISABLE CACHE TO FIX UNAVAILABLE MODEL ISSUE
    # Always create new instance to ensure current settings are used
    logger.info(f"Creating new tutor instance (cache disabled to fix model issue)")
    return UniversalSmartTutor(course_code, student_id)
